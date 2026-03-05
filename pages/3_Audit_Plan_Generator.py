import streamlit as st
from datetime import datetime
import re
from openai import OpenAI
import os, html, io, base64, json
import docx
from docx.enum.section import WD_ORIENT
from docx.shared import Pt
import streamlit.components.v1 as components

import sys, os, pathlib
_here = pathlib.Path(__file__).resolve().parent
for _p in [_here.parent, _here, pathlib.Path(os.getcwd())]:
    if (_p / "theme.py").exists():
        if str(_p) not in sys.path: sys.path.insert(0, str(_p))
        break
try:
    from theme import apply_theme, SIDEBAR_HTML
except ImportError:
    def apply_theme(): pass
    SIDEBAR_HTML = "<p style=\'color:white\'>AIT</p>"

st.set_page_config(layout="wide", page_title="Audit Plan Generator")
apply_theme()

with st.sidebar:
    st.markdown(SIDEBAR_HTML, unsafe_allow_html=True)

st.title("🔮 Audit Plan Generator")
st.markdown("เครื่องมือช่วยสร้างแผนและแนวการตรวจสอบ พร้อมระบบ AI ช่วยร่างเนื้อหา")

# ── State & Helpers ───────────────────────────────────
def init_plan_state():
    ss = st.session_state
    if "plan_gen_data" not in ss:
        ss.plan_gen_data = {"general_info":{"office":"","topic":"","agency":"","ministry":""},"objectives":[],"estimates":{"cost":"","effort":""},"signatures":{"maker":{"name":"","position":"","date":None,"comment":""},"reviewer":{"name":"","position":"","date":None,"comment":""},"approver":{"name":"","position":"","date":None,"comment":""}}}
    if "ui_feedback_message" not in ss: ss.ui_feedback_message = None
    if 'api_key_global' not in ss:
        try: ss['api_key_global'] = st.secrets["api_key"]
        except (KeyError, FileNotFoundError): ss['api_key_global'] = ""

init_plan_state()

def add_objective():
    new_obj = {"id":f"obj_{len(st.session_state.plan_gen_data['objectives'])+1}","text":"","issues":[]}
    st.session_state.plan_gen_data["objectives"].append(new_obj)
    st.session_state.ui_feedback_message = None

def remove_objective(obj_index):
    st.session_state.plan_gen_data["objectives"].pop(obj_index)
    st.session_state.ui_feedback_message = None

def add_issue(obj_index, parent_issue_path=None):
    obj = st.session_state.plan_gen_data["objectives"][obj_index]
    target_container = obj
    if parent_issue_path:
        for index in parent_issue_path: target_container = target_container["issues"][index]
    new_issue = {"id":f"issue_{obj_index}_{len(target_container['issues'])+1}","text":"","details":{"criteria":"","info_needed":"","source":"","collection_method":"","analysis_method":""},"issues":[]}
    target_container["issues"].append(new_issue)
    st.session_state.ui_feedback_message = None

def run_ai_for_field(obj_index, path, field_name):
    st.session_state.ui_feedback_message = None
    try:
        api_key = st.session_state.get("api_key_global")
        if not api_key:
            st.session_state.ui_feedback_message = ("error","ไม่พบ API Key"); return
        client = OpenAI(api_key=api_key, base_url="https://api.opentyphoon.ai/v1")
        obj = st.session_state.plan_gen_data["objectives"][obj_index]
        target_issue = obj
        for index in path: target_issue = target_issue["issues"][index]
        info = st.session_state.plan_gen_data['general_info']
        context = f"เรื่องที่ตรวจสอบ: {info.get('topic','')} (หน่วยงาน: {info.get('agency','')}, กระทรวง: {info.get('ministry','')})\n"
        context += f"วัตถุประสงค์: {obj.get('text','')}\nประเด็น: {target_issue.get('text','')}\n"
        instructions = {
            "criteria":          "จงสร้างเฉพาะ 'เกณฑ์การตรวจสอบ' (Audit Criteria) ที่เหมาะสม",
            "info_needed":       "จงระบุ 'ข้อมูลที่ต้องการ' เพื่อสนับสนุนข้อตรวจพบ สรุปผล และข้อเสนอแนะ",
            "source":            "จงระบุ 'แหล่งข้อมูล' ที่จะสามารถรวบรวมข้อมูลได้",
            "collection_method": "จงระบุ 'วิธีการรวบรวมหลักฐาน' เช่น การสุ่มตัวอย่าง การตรวจสอบเอกสาร การสัมภาษณ์",
            "analysis_method":   "จงระบุ 'วิธีการวิเคราะห์หลักฐาน' ที่จะใช้ในการประมวลผล",
        }
        full_prompt = f"คุณคือผู้เชี่ยวชาญด้านการตรวจสอบภาครัฐ Performance Audit\n{context}\n**คำสั่ง:** {instructions.get(field_name,'')}\nตอบเป็นข้อความธรรมดา bullet points เท่านั้น"
        response = client.chat.completions.create(model="typhoon-v2.5-30b-a3b-instruct", messages=[{"role":"user","content":full_prompt}], temperature=0.5, max_tokens=4096)
        cleaned_text = response.choices[0].message.content.strip().replace("**","")
        if cleaned_text:
            target_issue['details'][field_name] = cleaned_text
            key_suffix = f"{obj_index}_{'_'.join(map(str,path))}"
            st.session_state[f"{field_name}_{key_suffix}"] = cleaned_text
            st.session_state.ui_feedback_message = ("success", f"AI สร้าง '{field_name}' เรียบร้อยแล้ว")
        else:
            st.session_state.ui_feedback_message = ("error", f"AI ไม่สามารถสร้างเนื้อหาสำหรับ '{field_name}'")
    except Exception as e:
        st.session_state.ui_feedback_message = ("error", f"เกิดข้อผิดพลาด: {e}")

@st.cache_data
def load_font_as_base64(font_path):
    try:
        with open(font_path,"rb") as f: return base64.b64encode(f.read()).decode()
    except FileNotFoundError: return None

def generate_html_report(data):
    sarabun_regular_b64 = load_font_as_base64("Sarabun-Regular.ttf")
    sarabun_bold_b64    = load_font_as_base64("Sarabun-Bold.ttf")
    font_faces = ""
    if sarabun_regular_b64: font_faces += f"@font-face {{font-family:'Sarabun';src:url(data:font/truetype;charset=utf-8;base64,{sarabun_regular_b64}) format('truetype');font-weight:normal;}}\n"
    if sarabun_bold_b64:    font_faces += f"@font-face {{font-family:'Sarabun';src:url(data:font/truetype;charset=utf-8;base64,{sarabun_bold_b64}) format('truetype');font-weight:bold;}}\n"
    def fmt(text): return html.escape(text or "").replace('\n','<br>')
    html_template = f"""<!DOCTYPE html><html lang="th"><head><meta charset="UTF-8"><title>แผนและแนวการตรวจสอบ</title>
    <style>{font_faces}body{{font-family:'Sarabun',sans-serif;font-size:16px;margin:1cm;}}h2{{text-align:center;font-weight:bold;}}table{{width:100%;border-collapse:collapse;margin-top:1em;margin-bottom:1em;}}th,td{{border:1px solid black;padding:8px;text-align:left;vertical-align:top;}}.signature-table td{{height:120px;}}.print-button-container{{text-align:center;margin:15px 0;}}.print-button{{padding:10px 20px;font-size:16px;cursor:pointer;border-radius:5px;border:1px solid #7A2020;background-color:#f8e8e8;color:#7A2020;font-family:'Sarabun',sans-serif;}}thead th{{position:-webkit-sticky;position:sticky;top:-1px;background-color:#f8f9ee;z-index:2;}}@media print{{.no-print{{display:none;}}@page{{size:A4 landscape;margin:1.5cm;}}body{{margin:0;}}}}</style></head><body>
    <div class="print-button-container no-print"><button class="print-button" onclick="window.print()">🖨️ พิมพ์ / บันทึกเป็น PDF</button></div>
    <h2>แผนและแนวการตรวจสอบ</h2>
    <p><b>สำนักงาน/จังหวัด:</b> {fmt(data['general_info'].get('office',''))}&nbsp;&nbsp;<b>เรื่องที่ตรวจสอบ:</b> {fmt(data['general_info'].get('topic',''))}</p>
    <p><b>หน่วยงาน:</b> {fmt(data['general_info'].get('agency',''))}&nbsp;&nbsp;<b>กระทรวง:</b> {fmt(data['general_info'].get('ministry',''))}</p>
    {''.join([f"<p><b>วัตถุประสงค์การตรวจสอบที่ {i+1}:</b> {fmt(obj.get('text',''))}</p>{''.join([f'<table><thead><tr><th>เกณฑ์การตรวจสอบ</th><th>ข้อมูลที่ต้องการ</th><th>แหล่งข้อมูล</th><th>วิธีรวบรวมหลักฐาน</th><th>วิธีวิเคราะห์หลักฐาน</th></tr></thead><tbody><tr><td colspan=\"5\"><b>ประเด็น {i+1}.{j+1}:</b> {fmt(issue.get(\"text\",\"\"))}</td></tr><tr><td>{fmt(issue[\"details\"].get(\"criteria\",\"\"))}</td><td>{fmt(issue[\"details\"].get(\"info_needed\",\"\"))}</td><td>{fmt(issue[\"details\"].get(\"source\",\"\"))}</td><td>{fmt(issue[\"details\"].get(\"collection_method\",\"\"))}</td><td>{fmt(issue[\"details\"].get(\"analysis_method\",\"\"))}</td></tr></tbody></table>' for j,issue in enumerate(obj.get('issues',[]))])}" for i,obj in enumerate(data['objectives'])])}
    <p><b>ประมาณการค่าใช้จ่าย:</b> {fmt(data['estimates'].get('cost',''))}</p>
    <p><b>ประมาณการคน/วัน:</b> {fmt(data['estimates'].get('effort',''))}</p>
    <table class="signature-table"><thead><tr style="font-weight:bold;text-align:center;"><th>ผู้จัดทำ</th><th>ผู้สอบทาน</th><th>ผู้อนุมัติ</th></tr></thead><tbody><tr>
    <td><b>ลงชื่อ:</b> {fmt(data['signatures']['maker'].get('name',''))}<br><b>ตำแหน่ง:</b> {fmt(data['signatures']['maker'].get('position',''))}<br><b>วันที่:</b> {data['signatures']['maker'].get('date').strftime('%d/%m/%Y') if data['signatures']['maker'].get('date') else ''}</td>
    <td><b>ลงชื่อ:</b> {fmt(data['signatures']['reviewer'].get('name',''))}<br><b>ตำแหน่ง:</b> {fmt(data['signatures']['reviewer'].get('position',''))}<br><b>วันที่:</b> {data['signatures']['reviewer'].get('date').strftime('%d/%m/%Y') if data['signatures']['reviewer'].get('date') else ''}</td>
    <td><b>ลงชื่อ:</b> {fmt(data['signatures']['approver'].get('name',''))}<br><b>ตำแหน่ง:</b> {fmt(data['signatures']['approver'].get('position',''))}<br><b>วันที่:</b> {data['signatures']['approver'].get('date').strftime('%d/%m/%Y') if data['signatures']['approver'].get('date') else ''}</td>
    </tr></tbody></table></body></html>"""
    return html_template

def generate_docx_report(data):
    doc = docx.Document()
    sec = doc.sections[-1]; new_w, new_h = sec.page_height, sec.page_width
    sec.orientation = WD_ORIENT.LANDSCAPE; sec.page_width = new_w; sec.page_height = new_h
    font = doc.styles['Normal'].font; font.name = 'TH SarabunPSK'; font.size = Pt(14)
    doc.add_heading('แผนและแนวการตรวจสอบ', level=1)
    info = data["general_info"]
    doc.add_paragraph(f"เรื่องที่ตรวจสอบ: {info.get('topic','N/A')}    หน่วยงาน: {info.get('agency','N/A')}    กระทรวง: {info.get('ministry','N/A')}")
    doc.add_paragraph(f"สำนักงาน/จังหวัด: {info.get('office','N/A')}")
    for i, obj in enumerate(data["objectives"]):
        doc.add_paragraph().add_run(f"วัตถุประสงค์ที่ {i+1}: {obj.get('text','')}").bold = True
        for j, issue in enumerate(obj.get('issues',[])):
            doc.add_paragraph(f"ประเด็น {i+1}.{j+1}: {issue.get('text','')}")
            if not issue.get('issues'):
                d = issue.get('details',{}); tbl = doc.add_table(rows=1, cols=5); tbl.style = 'Table Grid'
                hdr = tbl.rows[0].cells
                hdr[0].text='เกณฑ์'; hdr[1].text='ข้อมูลที่ต้องการ'; hdr[2].text='แหล่งข้อมูล'; hdr[3].text='วิธีรวบรวม'; hdr[4].text='วิธีวิเคราะห์'
                row = tbl.add_row().cells
                row[0].text=d.get('criteria',''); row[1].text=d.get('info_needed',''); row[2].text=d.get('source',''); row[3].text=d.get('collection_method',''); row[4].text=d.get('analysis_method','')
                doc.add_paragraph()
    doc.add_paragraph(f"ประมาณการค่าใช้จ่าย: {data['estimates'].get('cost','')}"); doc.add_paragraph(f"ประมาณการคน/วัน: {data['estimates'].get('effort','')}")
    buffer = io.BytesIO(); doc.save(buffer); buffer.seek(0)
    return buffer

# ── Feedback ──────────────────────────────────────────
if st.session_state.get("ui_feedback_message"):
    msg_type, msg_content = st.session_state.ui_feedback_message
    if msg_type == "success": st.success(msg_content)
    else: st.error(msg_content)
    st.session_state.ui_feedback_message = None

# ── Section 1: General Info ───────────────────────────
with st.form("general_info_form"):
    st.subheader("1. ข้อมูลทั่วไป")
    c1, c2 = st.columns(2)
    st.session_state.plan_gen_data["general_info"]["office"]  = c1.text_input("สำนักงาน/จังหวัด/กลุ่ม", st.session_state.plan_gen_data["general_info"]["office"])
    st.session_state.plan_gen_data["general_info"]["topic"]   = c1.text_input("เรื่องที่ตรวจสอบ", st.session_state.plan_gen_data["general_info"]["topic"])
    st.session_state.plan_gen_data["general_info"]["agency"]  = c2.text_input("หน่วยงาน", st.session_state.plan_gen_data["general_info"]["agency"])
    st.session_state.plan_gen_data["general_info"]["ministry"]= c2.text_input("กระทรวง", st.session_state.plan_gen_data["general_info"]["ministry"])
    st.form_submit_button("💾 บันทึกข้อมูลทั่วไป", use_container_width=True)

# ── Section 2: Objectives & Issues ───────────────────
st.subheader("2. วัตถุประสงค์และประเด็นการตรวจสอบ")
st.write("ระบุวัตถุประสงค์และเพิ่มประเด็นการตรวจสอบย่อย เพื่อให้ AI สร้างรายละเอียดได้อย่างแม่นยำ")

for i, obj in enumerate(st.session_state.plan_gen_data["objectives"]):
    with st.container(border=True):
        c1, c2 = st.columns([5,1])
        st.session_state.plan_gen_data["objectives"][i]['text'] = c1.text_area(f"วัตถุประสงค์ที่ {i+1}", obj.get('text',''), key=f"obj_text_{i}")
        c2.button("🗑️ ลบ", key=f"del_obj_{i}", on_click=remove_objective, args=(i,), use_container_width=True)

        def display_issues(issues_list, obj_index, path):
            for j, issue in enumerate(issues_list):
                current_path = path + [j]
                prefix = ".".join([str(obj_index+1)] + [str(p+1) for p in current_path])
                key_suffix = f"{obj_index}_{'_'.join(map(str,current_path))}"
                with st.container():
                    st.markdown(f"<div style='margin-left:{len(current_path)*20}px;'>", unsafe_allow_html=True)
                    target_issue = issue
                    target_issue['text'] = st.text_area(f"ประเด็นการตรวจสอบที่ {prefix}", value=target_issue.get('text',''), key=f"issue_text_{key_suffix}")
                    if not target_issue.get('issues'):
                        with st.expander("💡 เพิ่มรายละเอียดแนวการตรวจสอบ (AI ช่วยได้)"):
                            details = target_issue.get('details',{})
                            field_map = {"criteria":"เกณฑ์การตรวจสอบ","info_needed":"ข้อมูลที่ต้องการ","source":"แหล่งข้อมูล","collection_method":"วิธีการรวบรวมหลักฐาน","analysis_method":"วิธีการวิเคราะห์หลักฐาน"}
                            for field, label in field_map.items():
                                col1, col2 = st.columns([4,1])
                                with col1: details[field] = st.text_area(label, value=details.get(field,''), key=f"{field}_{key_suffix}")
                                with col2: st.button(f"✨ AI", key=f"ai_btn_{field}_{key_suffix}", on_click=run_ai_for_field, args=(i, current_path, field))
                    st.button(f"➕ เพิ่มประเด็นย่อย ({prefix})", key=f"add_sub_{key_suffix}", on_click=add_issue, args=(obj_index, current_path))
                    if target_issue.get('issues'): display_issues(target_issue['issues'], obj_index, current_path)
                    st.markdown("</div>", unsafe_allow_html=True)

        if obj.get('issues'): display_issues(obj['issues'], i, [])
        st.button("➕ เพิ่มประเด็นการตรวจสอบหลัก", key=f"add_issue_{i}", on_click=add_issue, args=(i, None))

st.button("➕ เพิ่มวัตถุประสงค์", on_click=add_objective, type="primary")

# ── Section 3: Estimates & Signatures ────────────────
with st.form("estimates_signatures_form"):
    st.subheader("3. ประมาณการและผู้จัดทำ")
    st.session_state.plan_gen_data["estimates"]["cost"]   = st.text_area("ประมาณการค่าใช้จ่ายในการตรวจสอบ", st.session_state.plan_gen_data["estimates"]["cost"])
    st.session_state.plan_gen_data["estimates"]["effort"] = st.text_area("ประมาณการคน/วันที่ใช้ในการตรวจสอบ", st.session_state.plan_gen_data["estimates"]["effort"])
    c1, c2, c3 = st.columns(3)
    sig_data = st.session_state.plan_gen_data["signatures"]
    with c1:
        st.markdown("**ผู้จัดทำ**")
        sig_data["maker"]["name"]     = st.text_input("ลงชื่อ", value=sig_data["maker"].get("name",""), key="maker_name")
        sig_data["maker"]["position"] = st.text_input("ตำแหน่ง", value=sig_data["maker"].get("position",""), key="maker_pos")
        sig_data["maker"]["date"]     = st.date_input("วันที่", value=sig_data["maker"].get("date"), key="maker_date")
        sig_data["maker"]["comment"]  = st.text_area("ความเห็นเพิ่มเติม", value=sig_data["maker"].get("comment",""), key="maker_comment")
    with c2:
        st.markdown("**ผู้สอบทาน**")
        sig_data["reviewer"]["name"]     = st.text_input("ลงชื่อ", value=sig_data["reviewer"].get("name",""), key="reviewer_name")
        sig_data["reviewer"]["position"] = st.text_input("ตำแหน่ง", value=sig_data["reviewer"].get("position",""), key="reviewer_pos")
        sig_data["reviewer"]["date"]     = st.date_input("วันที่", value=sig_data["reviewer"].get("date"), key="reviewer_date")
        sig_data["reviewer"]["comment"]  = st.text_area("ความเห็นเพิ่มเติม", value=sig_data["reviewer"].get("comment",""), key="reviewer_comment")
    with c3:
        st.markdown("**ผู้อนุมัติ (รผต. / ผอ. สำนัก)**")
        sig_data["approver"]["name"]     = st.text_input("ลงชื่อ", value=sig_data["approver"].get("name",""), key="approver_name")
        sig_data["approver"]["position"] = st.text_input("ตำแหน่ง", value=sig_data["approver"].get("position",""), key="approver_pos")
        sig_data["approver"]["date"]     = st.date_input("วันที่", value=sig_data["approver"].get("date"), key="approver_date")
        sig_data["approver"]["comment"]  = st.text_area("ความเห็นเพิ่มเติม", value=sig_data["approver"].get("comment",""), key="approver_comment")
    st.form_submit_button("💾 บันทึกข้อมูลผู้จัดทำ", use_container_width=True)

st.divider()

# ── Section 4: Preview & Export ───────────────────────
with st.container(border=True):
    st.subheader("4. แสดงผลและส่งออกเอกสาร")
    html_report = generate_html_report(st.session_state.plan_gen_data)
    components.html(html_report, height=800, scrolling=True)
    st.markdown("---")
    st.markdown("##### ดาวน์โหลดเป็นไฟล์ Word")
    docx_buffer = generate_docx_report(st.session_state.plan_gen_data)
    st.download_button(
        label="📂 ดาวน์โหลดเป็นไฟล์ Word (.docx)",
        data=docx_buffer,
        file_name=f"audit_plan_{datetime.now().strftime('%Y%m%d')}.docx",
        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        use_container_width=True
    )
